import streamlit as st
import pandas as pd
import io
import time
import hashlib
import random
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
import matplotlib.pyplot as plt
from PIL import Image

# Intentar leer la rúbrica desde RubricaFinal.docx si existe
try:
    from docx import Document
    def leer_rubrica_docx(path="RubricaFinal.docx"):
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return text
    rubrica_text = leer_rubrica_docx("RubricaFinal.docx")
except Exception:
    rubrica_text = None

# Rúbrica por defecto (si no se puede leer el docx)
DEFAULT_RUBRICA = {
    "Contexto y relevancia": {"max": 4},
    "Revisión de literatura": {"max": 4},
    "Identificación del problema": {"max": 4},
    "Objetivos/preguntas": {"max": 4},
    "Justificación y contribución": {"max": 2},
    "Estructura y fluidez": {"max": 2},
}
RUBRICA = DEFAULT_RUBRICA  # usa la integrada; si deseas parsear rubrica_text, puedes hacerlo luego

# --- CSS para mejor apariencia ---
st.set_page_config(page_title="Revisor Falso de Artículos", page_icon="📝", layout="wide")
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 1.4rem;
        border-radius: 10px;
        text-align: center;
        color: white;
        margin-bottom: 1.2rem;
    }
    .metric-card {
        background: white;
        padding: 1rem;
        border-radius: 8px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.06);
        border-left: 4px solid #667eea;
    }
    .processing-box {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-header"><h1>📝 Revisor Automático de Artículos</h1><p>Realiza una revisión basada en la rúbrica con IA y genera reportes y PDF</p></div>', unsafe_allow_html=True)

# Sidebar: info y subida de rúbrica opcional
with st.sidebar:
    st.header("⚙️ Configuración")
    st.write("Sube aquí (opcional) tu rúbrica en .docx si quieres usarla (se prioriza `RubricaFinal.docx` en el repo).")
    rub_upload = st.file_uploader("Subir rúbrica (.docx)", type=["docx"], accept_multiple_files=False)
    if rub_upload:
        try:
            from docx import Document
            doc = Document(rub_upload)
            rub_text_user = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            st.success("✅ Rúbrica cargada desde archivo.")
            rubrica_text = rub_text_user
            # Nota: para parseo avanzado puedes implementar aquí
        except Exception as e:
            st.warning("⚠️ No se pudo leer la rúbrica. Se usará la rúbrica por defecto.")
    st.markdown("---")
    st.info("""
    Instrucciones rápidas:
    - Sube hasta 30 archivos PDF (artículos).
    - Presiona **Evaluar**: el sistema simulará la revisión y generará notas y comentarios.
    - Puedes descargar un PDF con todos los resultados.
    """)
    st.markdown("---")
    st.write("Sistema revisor de articulos con IA.")

# ====================== INTERFAZ ======================
st.header("1️⃣ Subir artículos (PDF)")
st.info("Puedes subir hasta 30 artículos para que la IA evalue según la rúbrica.")
uploaded_files = st.file_uploader("Sube los PDFs de los artículos", type=["pdf"], accept_multiple_files=True)

# Límite 5
if uploaded_files and len(uploaded_files) > 5:
    st.error("Máximo 5 archivos permitidos. Reduce la cantidad.")
    uploaded_files = uploaded_files[:5]

# Mostrar rúbrica (resumida)
st.header("Rúbrica usada")
st.write("Se usarán los criterios y puntajes máximos siguientes (total = 20 puntos):")
rub_df = pd.DataFrame([{"Criterio": k, "Max Pts": v["max"]} for k, v in RUBRICA.items()])
st.table(rub_df.set_index("Criterio"))

# ====================== FUNCIONES ======================
def seed_from_name(name: str):
    h = hashlib.sha256(name.encode("utf-8")).hexdigest()
    seed = int(h[:16], 16) % (2**31)
    return seed

# Generador de comentarios por criterio (plantillas)
COMMENT_TEMPLATES = {
    "Contexto y relevancia": [
        "Contexto bien establecido; relevancia clara y bien argumentada.",
        "Buen contexto pero podría enfatizar más la contribución al área.",
        "Contexto limitado; falta justificación de por qué el problema es relevante.",
        "Contexto pobre o ausente; no queda claro por qué investigar esto."
    ],
    "Revisión de literatura": [
        "Revisión completa y crítica; referencias pertinentes y bien integradas.",
        "Buena revisión pero falta profundidad crítica en algunas referencias clave.",
        "Revisión superficial; faltan conexiones claras con el problema.",
        "Revisión insuficiente o con referencias irrelevantes."
    ],
    "Identificación del problema": [
        "Problema claramente identificado y bien fundamentado en la literatura.",
        "Problema identificado, pero requeriría mayor precisión en su delimitación.",
        "Problema poco definido o no claramente derivado de la literatura.",
        "No se identifica un problema claro."
    ],
    "Objetivos/preguntas": [
        "Objetivos claros, específicos y alineados con el problema.",
        "Objetivos aceptables pero podrían ser más medibles o precisos.",
        "Objetivos vagos o demasiado amplios.",
        "Objetivos confusos, ausentes o no evaluables."
    ],
    "Justificación y contribución": [
        "Justificación sólida; contribución teórica/práctica bien explicada.",
        "Justificación adecuada pero el impacto podría desarrollarse más.",
        "Justificación débil; contribuciones poco claras.",
        "No justifican la investigación ni la contribución."
    ],
    "Estructura y fluidez": [
        "Estructura lógica y flujo excelente; redacción académica clara.",
        "Buena estructura con algunas transiciones mejorables.",
        "Estructura desorganizada que afecta la comprensión.",
        "Estructura deficiente; difícil de seguir."
    ]
}

def evaluar_articulo_fake(file_like, rubrica=RUBRICA):
    """
    Genera una evaluación 'fake' reproducible basada en el nombre del archivo.
    Retorna dict con puntajes por criterio, comentarios y nota total.
    """
    nombre = getattr(file_like, "name", f"art_{time.time()}")
    seed = seed_from_name(nombre)
    rnd = random.Random(seed)
    
    detalle = {}
    total_obtenido = 0
    for criterio, meta in rubrica.items():
        max_pts = meta["max"]
        # Generar puntuación entera entre 0 y max_pts
        # Tendencia: la mayoría de artículos estén entre 50% y 95% del max, pero con variabilidad
        base = rnd.normalvariate(0.75 * max_pts, 0.9)
        pts = int(max(0, min(max_pts, round(base))))
        # Ajuste aleatorio extra
        if rnd.random() < 0.08:
            pts = max(0, pts - rnd.randint(1, max(1, max_pts//2)))
        if rnd.random() < 0.06:
            pts = min(max_pts, pts + rnd.randint(1, max(1, max_pts//2)))
        comentario = rnd.choice(COMMENT_TEMPLATES.get(criterio, ["Comentario genérico."]))
        # si la puntuación es baja, elegir plantilla más crítica
        if pts >= 0.9 * max_pts:
            comentario = COMMENT_TEMPLATES[criterio][0]
        elif pts >= 0.6 * max_pts:
            comentario = COMMENT_TEMPLATES[criterio][1]
        elif pts >= 0.3 * max_pts:
            comentario = COMMENT_TEMPLATES[criterio][2]
        else:
            comentario = COMMENT_TEMPLATES[criterio][3]
        
        detalle[criterio] = {"pts": pts, "max": max_pts, "comentario": comentario}
        total_obtenido += pts
    
    # Escala total ya es sobre 20 si rubrica suma 20
    nota = round(float(total_obtenido), 2)
    
    return {
        "nombre_pdf": nombre,
        "detalle": detalle,
        "total": nota
    }

def generar_reporte_pdf(buffer_io, resultados, curso_nombre="Revisión Artículos", curso_codigo="ART-REV"):
    """
    Genera un PDF en buffer_io (BytesIO) con resultados (lista de dicts generados arriba)
    """
    doc = SimpleDocTemplate(buffer_io, pagesize=A4, topMargin=2*cm, bottomMargin=2*cm)
    elementos = []
    styles = getSampleStyleSheet()
    titulo_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#667eea'),
        alignment=TA_CENTER
    )
    elementos.append(Paragraph("📝 REPORTE DE REVISIÓN SIMULADA", titulo_style))
    elementos.append(Spacer(1, 0.2*cm))
    elementos.append(Paragraph(f"Curso / Proyecto: {curso_nombre} — {curso_codigo}", styles['Normal']))
    elementos.append(Paragraph(f"Fecha: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}", styles['Normal']))
    elementos.append(Spacer(1, 0.5*cm))
    
    # Resumen tablas
    tabla_res = [['#', 'Nombre', 'Nota (0-20)']]
    for idx, r in enumerate(resultados, 1):
        tabla_res.append([str(idx), r['nombre_pdf'], f"{r['total']:.2f}"])
    t = Table(tabla_res, colWidths=[1.2*cm, 12*cm, 3*cm])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#667eea')),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.grey)
    ]))
    elementos.append(t)
    elementos.append(Spacer(1, 0.5*cm))
    
    # Detalle por artículo
    for idx, r in enumerate(resultados, 1):
        elementos.append(Paragraph(f"<b>{idx}. {r['nombre_pdf']} — Nota: {r['total']:.2f}</b>", styles['Heading3']))
        data_det = [['Criterio', 'Pts', 'Max', 'Comentario']]
        for crit, info in r['detalle'].items():
            data_det.append([crit, str(info['pts']), str(info['max']), info['comentario']])
        td = Table(data_det, colWidths=[6*cm, 1.5*cm, 1.5*cm, 7*cm])
        td.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f0f0f0')),
            ('ALIGN', (1,1), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('GRID', (0,0), (-1,-1), 0.3, colors.grey)
        ]))
        elementos.append(td)
        elementos.append(Spacer(1, 0.3*cm))
    
    doc.build(elementos)

# ====================== BOTÓN EVALUAR ======================
st.header("2️⃣ Evaluar")
col1, col2 = st.columns([2,1])
with col1:
    curso_nombre = st.text_input("Nombre del proyecto / curso", value="Revisión de Artículos")
with col2:
    curso_codigo = st.text_input("Código (opcional)", value="ART-REV")

if st.button("🚀 Evaluar artículos", disabled=(not uploaded_files)):
    if not uploaded_files:
        st.warning("Sube al menos 1 PDF para evaluar.")
    else:
        st.info("Iniciando evaluación simulada...")
        progreso = st.progress(0)
        resultados = []
        total = len(uploaded_files)
        for i, f in enumerate(uploaded_files):
            progreso.progress((i+1)/total)
            time.sleep(0.6)  # efecto visual
            r = evaluar_articulo_fake(f, RUBRICA)
            resultados.append(r)
        st.success("✅ Evaluación completada.")
        st.session_state.resultados = resultados
        st.session_state.curso_nombre = curso_nombre
        st.session_state.curso_codigo = curso_codigo

# ====================== MOSTRAR RESULTADOS ======================
if 'resultados' in st.session_state and st.session_state.resultados:
    resultados = st.session_state.resultados
    st.markdown("---")
    st.header("3️⃣ Resultados y estadísticas")
    
    df = pd.DataFrame([{
        "nombre_pdf": r['nombre_pdf'],
        "nota": r['total']
    } for r in resultados])
    
    # Métricas
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("📈 Promedio general", f"{df['nota'].mean():.2f}")
    with col2:
        aprobados = len(df[df['nota'] >= 14])
        st.metric("✅ Aprobados", f"{aprobados}", delta=f"{(aprobados/len(df)*100):.1f}%")
    with col3:
        st.metric("🏆 Nota más alta", f"{df['nota'].max():.2f}")
    with col4:
        st.metric("📉 Nota más baja", f"{df['nota'].min():.2f}")
    
    st.markdown("### 📋 Detalle de calificaciones")
    df_display = df.copy().sort_values('nota', ascending=False).reset_index(drop=True)
    df_display.index += 1
    st.dataframe(df_display.rename(columns={"nombre_pdf":"Nombre del Archivo","nota":"Nota (0-20)"}), use_container_width=True)
    
    # Gráficas: histograma y barras
    st.markdown("### 📊 Gráficas")
    fig1, ax1 = plt.subplots()
    ax1.hist(df['nota'], bins=5)
    ax1.set_title("Distribución de notas")
    ax1.set_xlabel("Nota (0-20)")
    ax1.set_ylabel("Cantidad")
    st.pyplot(fig1)
    
    fig2, ax2 = plt.subplots()
    ax2.barh(df['nombre_pdf'], df['nota'])
    ax2.set_xlabel("Nota")
    ax2.set_title("Notas por artículo")
    st.pyplot(fig2)
    
    st.markdown("---")
    st.header("4️⃣ Comentarios por artículo")
    for r in resultados:
        with st.expander(f"{r['nombre_pdf']} — Nota: {r['total']:.2f}"):
            for crit, info in r['detalle'].items():
                st.markdown(f"**{crit}** — {info['pts']}/{info['max']}")
                st.markdown(f"> {info['comentario']}")
    
    # Generar PDF completo
    st.markdown("---")
    st.header("5️⃣ Exportar reporte PDF")
    if st.button("📄 Generar y descargar PDF de resultados"):
        buffer = io.BytesIO()
        generar_reporte_pdf(buffer, resultados, st.session_state.curso_nombre, st.session_state.curso_codigo)
        buffer.seek(0)
        fn = f"reporte_revision_{st.session_state.curso_codigo}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        st.download_button("⬇️ Descargar PDF", data=buffer, file_name=fn, mime="application/pdf")

